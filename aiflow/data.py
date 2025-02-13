import json
import re
import os
import torch
from pydub import AudioSegment, silence
from transformers import pipeline
from tqdm import tqdm

class AudioDataWorker:
    def __init__(self, device="mps", model="openai/whisper-medium"):
        """Initialize the AudioDataWorker with a device and ASR model.

        Args:
            device (str, optional): Compute device (e.g. "mps"). Defaults to "mps".
            model (str, optional): Hugging Face model identifier. Defaults to "openai/whisper-medium".
        """
        self.device = device
        self.model = model
        self.asr_pipeline = pipeline(
            "automatic-speech-recognition",
            model=self.model,
            torch_dtype=torch.float32,
            device=self.device,
            model_kwargs={"attn_implementation": "sdpa"},
        )

    def combine_audio_files(self, input_folder, output_file, silence_duration=1000):
        """Combine .wav files in numeric order with a silence gap between them.

        Args:
            input_folder (str): Folder containing audio files.
            output_file (str): Path to save the combined audio.
            silence_duration (int, optional): Duration of silence (ms) between files. Defaults to 1000.
        """
        files = sorted(
            [f for f in os.listdir(input_folder) if f.endswith('.wav')],
            key=lambda x: int(os.path.splitext(x)[0])
        )
        combined = AudioSegment.empty()
        silence_seg = AudioSegment.silent(duration=silence_duration)
        for i, file in enumerate(files):
            audio = AudioSegment.from_wav(os.path.join(input_folder, file))
            combined += audio + (silence_seg if i < len(files) - 1 else AudioSegment.empty())
        combined.export(output_file, format="wav")
        print(f"Combined audio saved as {output_file}")

    def split_audio_into_chunks(self, input_file, output_dir, min_silence_len=400, silence_thresh=-40, min_chunk_length=4000, max_chunk_length=20000):
        """Split an audio file into chunks based on silence and length thresholds.

        Args:
            input_file (str): Path to input .wav file.
            output_dir (str): Directory to save the chunks.
            min_silence_len (int, optional): Minimum silence length (ms) to split on. Defaults to 400.
            silence_thresh (int, optional): Silence threshold (dBFS). Defaults to -40.
            min_chunk_length (int, optional): Minimum duration (ms) of a chunk. Defaults to 4000.
            max_chunk_length (int, optional): Maximum duration (ms) of a chunk. Defaults to 20000.
        """
        os.makedirs(output_dir, exist_ok=True)
        chunks = silence.split_on_silence(
            AudioSegment.from_wav(input_file),
            min_silence_len=min_silence_len,
            silence_thresh=silence_thresh,
            keep_silence=500
        )
        processed = []
        current = AudioSegment.empty()
        
        for chunk in chunks:
            current += chunk
            while len(current) >= min_chunk_length:
                split_point = min(len(current), max_chunk_length)
                processed.append(current[:split_point])
                current = current[split_point:] if split_point < len(current) else AudioSegment.empty()
                if len(current) < min_chunk_length: break
        
        if len(current) > 0: processed.append(current)
        for i, chunk in enumerate(processed, 1):
            chunk.export(f"{output_dir}/chunk_{i}.wav", format="wav")
        print("Splitting completed successfully!")

    def calculate_total_duration(self, directory):
        """Calculate and display the total duration of .wav files in a directory.

        Args:
            directory (str): Directory containing .wav files.
        """
        wav_files = [os.path.join(directory, f) for f in os.listdir(directory) if f.endswith('.wav')]
        total_ms = 0
        print("Calculating total duration...")
        for wav_file in tqdm(wav_files, desc="Processing files"):
            try:
                total_ms += len(AudioSegment.from_wav(wav_file))
            except Exception as e:
                print(f"Error processing {wav_file}: {e}")
        total_seconds = total_ms / 1000
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        print(f"\nTotal duration:\nHours: {int(hours)}\nMinutes: {int(minutes)}\nSeconds: {int(seconds)}")
        print(f"Total files processed: {len(wav_files)}")

    def transcribe_audio(self, audio_file, chunk_length_s=30, batch_size=24, return_timestamps=False):
        """Transcribe an audio file using the ASR pipeline.

        Args:
            audio_file (str): Audio file path.
            chunk_length_s (int, optional): Time (s) to split audio into chunks. Defaults to 30.
            batch_size (int, optional): Batch size for ASR processing. Defaults to 24.
            return_timestamps (bool, optional): Whether to return timestamps. Defaults to False.

        Returns:
            dict: Transcription result.
        """
        return self.asr_pipeline(
            audio_file,
            chunk_length_s=chunk_length_s,
            batch_size=batch_size,
            return_timestamps=return_timestamps,
        )

    def transcribe_directory(self, directory_path, output_file, chunk_length_s=30, batch_size=24, return_timestamps=False):
        """Transcribe all .wav files in a directory and save results to an output file.

        Args:
            directory_path (str): Folder containing audio files.
            output_file (str): File path to save transcriptions.
            chunk_length_s (int, optional): Chunk length in seconds for transcription. Defaults to 30.
            batch_size (int, optional): Batch size for transcription. Defaults to 24.
            return_timestamps (bool, optional): Whether to return timestamps. Defaults to False.
        """
        with open(output_file, 'w') as f:
            for filename in os.listdir(directory_path):
                if filename.endswith(".wav"):
                    audio_path = os.path.join(directory_path, filename)
                    transcription = self.transcribe_audio(audio_path, chunk_length_s, batch_size, return_timestamps)
                    f.write(f"{audio_path}|Yuna|EN|{transcription['text'].strip()}\n")
                    print(f"Transcribed {audio_path}")

    def delete_unlisted_files(self, metadata_file, directory):
        """Delete .wav files not listed in the metadata file.

        Args:
            metadata_file (str): Path to the metadata file.
            directory (str): Directory containing .wav files.
        """
        with open(metadata_file, 'r') as file:
            metadata_files = {line.split('|')[0].strip() for line in file if line.strip()}
        all_files = {os.path.join(directory, f) for f in os.listdir(directory) if f.endswith('.wav')}
        for file_path in all_files - metadata_files:
            os.remove(file_path)
            print(f"Deleted {file_path}")
        print("Deletion of unlisted files completed successfully!")

    def rename_files(self, metadata_file, directory, start_number=2565):
        """Sequentially rename files in metadata and update the metadata file.

        Args:
            metadata_file (str): Path to the metadata file.
            directory (str): Directory where files reside.
            start_number (int, optional): Starting number for renaming. Defaults to 2565.
        """
        with open(metadata_file, 'r') as file:
            metadata = [line.strip().split('|') for line in file if line.strip()]
        for i, data in enumerate(metadata):
            old_path = data[0]
            new_path = os.path.join(directory, f"{start_number + i}.wav")
            os.rename(old_path, new_path)
            data[0] = new_path
            print(f"Renamed {old_path} to {new_path}")
        with open(metadata_file, 'w') as file:
            for data in metadata:
                file.write('|'.join(data) + '\n')
        print("Renaming and metadata update completed successfully!")

    def sort_transcriptions(self, input_file, output_file):
        """Sort transcription lines based on the file number extracted from their path.

        Args:
            input_file (str): Path to the unsorted transcription file.
            output_file (str): Path to save sorted transcriptions.
        """
        with open(input_file, 'r') as file:
            lines = file.readlines()
        sorted_lines = sorted(lines, key=lambda x: int(x.split('/')[1].split('.')[0]))
        with open(output_file, 'w') as file:
            file.writelines(sorted_lines)

    def transliterate_file(self, input_path, output_path):
        """Transliterate the text (4th field) in each line of a file.

        Args:
            input_path (str): Path to the input file.
            output_path (str): Path to save the transliterated output.
        """
        from yuna.text.cyrillic import transliterate
        with open(input_path, 'r', encoding='utf-8') as infile, open(output_path, 'w', encoding='utf-8') as outfile:
            for line in infile:
                parts = line.split('|')
                if len(parts) >= 4:
                    parts[3] = transliterate(parts[3], source='ru', target='en')
                outfile.write('|'.join(parts))

    def start_ui(self):
        """Launch the AudioDataWorker UI."""
        transcript = input("Enter the path to the transcript file: ")
        from aiflow.AudioDataWorkerUI import main
        main(transcript)

class TextDataWorker:
    """A class for processing and converting text data between different formats."""

    def __init__(self, device="mps"):
        """Initialize TextDataWorker.

        Args:
            device (str, optional): Compute device. Defaults to "mps".
        """
        self.device = device
        self.model_path = "/Users/yuki/Documents/Github/yuna-ai/lib/models/yuna/yuna-ai-v4-q6_k.gguf"

    def split_text_to_jsonl(self, input_file, output_file, mode="split_by_block", block="<SPLITTEXT>", chunk_size=16384):
        """Split text into JSONL format using blocks or tokens.

        Args:
            input_file (str): Path to input text file.
            output_file (str): Path to output JSONL file.
            mode (str, optional): Split mode ('split_by_block' or 'split_by_token_chunks'). Defaults to "split_by_block".
            block (str, optional): Block separator. Defaults to "<SPLITTEXT>".
            chunk_size (int, optional): Token chunk size. Defaults to 16384.
        """
        from llama_cpp import Llama
        llm = Llama(model_path=self.model_path)
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()

        with open(output_file, 'w', encoding='utf-8') as f:
            if mode == "split_by_block":
                [f.write(json.dumps({"text": chunk.strip()}) + '\n') for chunk in re.split(block, content)]
            else:
                tokens = llm.tokenize(content.encode('utf-8'), special=True, add_bos=False)
                [f.write(json.dumps({"text": llm.detokenize(tokens[i:i+chunk_size], special=True).decode('utf-8')}) + '\n')
                 for i in range(0, len(tokens), chunk_size)]

    def count_tokens(self, input_file):
        """Count tokens in a text file using llama tokenizer.

        Args:
            input_file (str): Path to input text file.
        """
        from llama_cpp import Llama
        with open(input_file, 'r', encoding='utf-8') as f:
            print(f"Total tokens: {len(Llama(model_path=self.model_path).tokenize(f.read().encode('utf-8')))}")

    def convert_jsonl_to_text(self, input_file, output_file):
        """Convert JSONL file to plain text.

        Args:
            input_file (str): Path to input JSONL file.
            output_file (str): Path to output text file.
        """
        with open(input_file, 'r', encoding='utf-8') as f_in, open(output_file, 'w', encoding='utf-8') as f_out:
            [f_out.write(json.loads(line)['text'] + '\n') for line in f_in]

    def convert_docx_to_text(self, input_file, output_file):
        """Convert DOCX to text, preserving italics as markdown.

        Args:
            input_file (str): Path to input DOCX file.
            output_file (str): Path to output text file.
        """
        from docx import Document
        doc = Document(input_file)
        with open(output_file, 'w') as f:
            f.write(''.join([f"**{run.text}**" if run.italic else run.text 
                           for para in doc.paragraphs 
                           for run in para.runs] + ['\n' for _ in doc.paragraphs]))